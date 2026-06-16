import os
import io
import uuid
import datetime
from typing import List, Optional
from fastapi import FastAPI, Depends, HTTPException, Header, UploadFile, File, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv

load_dotenv()

from database import SupabaseDB
from rag import EnhancedRAG
from utils import scrape_webpage

# Initialize FastAPI app
app = FastAPI(title="AskMyDocs API", version="1.0.0")

# Enable CORS for local Vite-React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Supabase Database
supabase_url = os.environ.get("SUPABASE_URL")
supabase_key = os.environ.get("SUPABASE_KEY")
if not supabase_url or not supabase_key:
    print("Warning: SUPABASE_URL or SUPABASE_KEY not set in environment.")

db = SupabaseDB(supabase_url, supabase_key)

# Request Models
class UserSignup(BaseModel):
    email: str
    password: str
    name: str

class UserLogin(BaseModel):
    email: str
    password: str

class NotebookCreate(BaseModel):
    name: str
    description: Optional[str] = ""
    color: Optional[str] = "#1E88E5"
    domains: Optional[List[str]] = ["General"]

class NotebookUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    color: Optional[str] = None
    domains: Optional[List[str]] = None

class ChatRequest(BaseModel):
    query: str
    history: Optional[List[dict]] = []
    mode: str = "direct_retrieval"  # direct_retrieval, deep_search, hybrid
    notebook_id: Optional[str] = None
    llm_model: Optional[str] = "llama-3.3-70b-versatile"
    embedding_model: Optional[str] = "sentence-transformers/all-MiniLM-L6-v2"
    chunk_size: Optional[int] = 1000
    chunk_overlap: Optional[int] = 200
    use_gpu: Optional[bool] = True

class UrlUploadRequest(BaseModel):
    url: str
    use_rag: Optional[bool] = True
    llm_model: Optional[str] = "llama-3.3-70b-versatile"
    embedding_model: Optional[str] = "sentence-transformers/all-MiniLM-L6-v2"
    chunk_size: Optional[int] = 1000
    chunk_overlap: Optional[int] = 200
    use_gpu: Optional[bool] = True
    split_strategy: Optional[str] = "semantic"

# Helper class to wrap uploaded file bytes for LangChain processing
class UploadedFileWrapper:
    def __init__(self, filename: str, content: bytes):
        self.name = filename
        self.content = content
        self._io = io.BytesIO(content)
        
    def read(self, size=-1):
        return self._io.read(size)
        
    def seek(self, offset, whence=0):
        return self._io.seek(offset, whence)
        
    def tell(self):
        return self._io.tell()

# Authentication dependency
async def get_current_user(authorization: str = Header(None), token: Optional[str] = None):
    session_id = None
    if authorization and authorization.startswith("Bearer "):
        session_id = authorization.split(" ")[1]
    elif token:
        session_id = token
        
    if not session_id:
        raise HTTPException(status_code=401, detail="Missing or invalid token")
    
    success, user_data = db.validate_session(session_id)
    if not success:
        raise HTTPException(status_code=401, detail=f"Session invalid or expired: {user_data}")
        
    # user_data contains: user_id, name, email
    user_data["session_id"] = session_id
    return user_data

# Background task to compile RAG embeddings asynchronously
def build_rag_index_task(notebook_id: str, user_id: str, files_data: List[dict], llm_model: str, embedding_model: str, chunk_size: int, chunk_overlap: int, use_gpu: bool, split_strategy: str = "semantic"):
    try:
        # Reconstruct adapted files
        adapted_files = []
        for fd in files_data:
            adapted_files.append(UploadedFileWrapper(fd["filename"], fd["content"]))
            
        print(f"Starting RAG embedding build for notebook {notebook_id} with {len(adapted_files)} files...")
        rag = EnhancedRAG(
            llm_model_name=llm_model,
            embedding_model_name=embedding_model,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            use_gpu=use_gpu
        )
        
        # Load existing index if any
        rag.load_vector_store(db, notebook_id)
        
        # Process files
        success = rag.process_files(
            files=adapted_files,
            user_id=user_id,
            mongodb=db,
            notebook_id=notebook_id,
            is_nested=False,
            split_strategy=split_strategy
        )
        
        if success:
            # Save vector store to database
            rag.save_vector_store(db, notebook_id, user_id)
            print(f"RAG embedding index built successfully for notebook {notebook_id}.")
        else:
            print(f"Failed to process files inside build task for notebook {notebook_id}: {rag.errors}")
    except Exception as e:
        print(f"Error in background build task: {str(e)}")

def rebuild_notebook_index_task(notebook_id: str, user_id: str):
    try:
        # 1. Fetch remaining documents in database
        success, documents = db.list_user_documents(user_id, notebook_id)
        if not success:
            print(f"Failed to list remaining documents for index rebuild: {documents}")
            return
        
        # If no documents are left, delete the index record and storage files
        if not documents:
            print(f"No documents left for notebook {notebook_id}. Deleting index...")
            try:
                # Remove index records and storage files
                idx_res = db.client.table("faiss_indexes").select("*").eq("notebook_id", notebook_id).execute()
                if idx_res.data:
                    for idx in idx_res.data:
                        try:
                            db.client.storage.from_("faiss-indexes").remove([idx["faiss_index_path"]])
                            if idx.get("documents_path"):
                                db.client.storage.from_("faiss-indexes").remove([idx["documents_path"]])
                        except Exception as storage_err:
                            print(f"Error removing storage index files: {str(storage_err)}")
                    db.client.table("faiss_indexes").delete().eq("notebook_id", notebook_id).execute()
                print("Index deleted successfully.")
            except Exception as e:
                print(f"Error deleting empty index: {str(e)}")
            return
            
        # 2. Download all remaining documents from storage
        adapted_files = []
        for doc in documents:
            doc_success, doc_data = db.get_document_file(doc["file_id"])
            if doc_success:
                adapted_files.append(UploadedFileWrapper(doc_data["filename"], doc_data["data"]))
                
        if not adapted_files:
            print(f"No file bytes retrieved for rebuild of notebook {notebook_id}.")
            return
            
        print(f"Rebuilding RAG index for notebook {notebook_id} with {len(adapted_files)} files...")
        # 3. Create a clean RAG instance (without loading existing index)
        rag = EnhancedRAG(
            llm_model_name="llama-3.3-70b-versatile",
            embedding_model_name="sentence-transformers/all-MiniLM-L6-v2",
            use_gpu=True
        )
        
        # Process all files (creating a fresh vector store)
        success = rag.process_files(
            files=adapted_files,
            user_id=user_id,
            mongodb=db,
            notebook_id=notebook_id,
            is_nested=False,
            split_strategy="semantic"
        )
        
        if success:
            # Save the clean vector store
            rag.save_vector_store(db, notebook_id, user_id)
            print(f"Rebuild completed successfully for notebook {notebook_id}.")
        else:
            print(f"Rebuild failed processing files for notebook {notebook_id}: {rag.errors}")
    except Exception as e:
        print(f"Error rebuilding notebook index: {str(e)}")

# Helper to resolve or create the hidden "Global Workspace" notebook for a user
def get_or_create_global_notebook(user_id: str) -> str:
    success, notebooks = db.get_notebooks(user_id)
    if success:
        for nb in notebooks:
            if nb.get("name") == "Global Workspace":
                return str(nb["id"])
    
    # Create Global Workspace notebook if not exists
    success, result = db.create_notebook(
        user_id,
        "Global Workspace",
        "Global workspace for normal chat and documents",
        "#4F46E5",
        {"domains": ["General"]}
    )
    if success:
        return str(result)
    else:
        raise HTTPException(status_code=500, detail=f"Failed to create Global Workspace: {result}")

# Routes

@app.post("/api/auth/signup")
async def signup(user: UserSignup):
    success, result = db.create_user(user.email, user.password, user.name)
    if not success:
        raise HTTPException(status_code=400, detail=result)
    return {"message": "User created successfully", "user_id": result}

@app.post("/api/auth/login")
async def login(user: UserLogin):
    success, result = db.authenticate_user(user.email, user.password)
    if not success:
        raise HTTPException(status_code=401, detail=result)
    return result # Returns name, email, user_id, session_id

@app.get("/api/auth/session")
async def check_session(current_user = Depends(get_current_user)):
    return current_user

class ChangePasswordRequest(BaseModel):
    old_password: str
    new_password: str

@app.post("/api/auth/change-password")
async def change_password(request: ChangePasswordRequest, current_user = Depends(get_current_user)):
    success, message = db.change_password(current_user["user_id"], request.old_password, request.new_password)
    if not success:
        raise HTTPException(status_code=400, detail=message)
    return {"message": message}

@app.delete("/api/auth/delete-account")
async def delete_account(current_user = Depends(get_current_user)):
    success, message = db.delete_user_account(current_user["user_id"])
    if not success:
        raise HTTPException(status_code=400, detail=message)
    return {"message": message}

@app.post("/api/auth/logout")
async def logout(current_user = Depends(get_current_user)):
    success, message = db.logout_user(current_user["session_id"])
    if not success:
        raise HTTPException(status_code=400, detail=message)
    return {"message": message}

# Notebooks CRUD

@app.get("/api/notebooks")
async def list_notebooks(current_user = Depends(get_current_user)):
    success, notebooks = db.get_notebooks(current_user["user_id"])
    if not success:
        raise HTTPException(status_code=400, detail=notebooks)
    
    # Format notebooks dates
    formatted = []
    for nb in notebooks:
        if nb.get("name") == "Global Workspace":
            continue
        formatted.append({
            "id": nb["id"],
            "name": nb["name"],
            "description": nb.get("description", ""),
            "color": nb.get("color", "#1E88E5"),
            "domains": nb.get("domains", ["General"]),
            "is_favorite": nb.get("is_favorite", False),
            "document_count": nb.get("document_count", 0),
            "rag_document_count": nb.get("rag_document_count", 0),
            "created_at": nb["created_at"].isoformat() if isinstance(nb["created_at"], datetime.datetime) else nb["created_at"],
            "last_accessed": nb["last_accessed"].isoformat() if isinstance(nb["last_accessed"], datetime.datetime) else nb["last_accessed"]
        })
    return formatted

@app.post("/api/notebooks")
async def create_notebook(notebook: NotebookCreate, current_user = Depends(get_current_user)):
    success, result = db.create_notebook(
        current_user["user_id"],
        notebook.name,
        notebook.description,
        notebook.color,
        {"domains": notebook.domains}
    )
    if not success:
        raise HTTPException(status_code=400, detail=result)
    return {"id": result, "name": notebook.name}

@app.get("/api/notebooks/{notebook_id}")
async def get_notebook(notebook_id: str, current_user = Depends(get_current_user)):
    resolved_id = get_or_create_global_notebook(current_user["user_id"]) if notebook_id == "global" else notebook_id
    success, notebook = db.get_notebook(resolved_id)
    if not success:
        raise HTTPException(status_code=404, detail=notebook)
    
    # Verify owner
    if str(notebook["user_id"]) != current_user["user_id"]:
        raise HTTPException(status_code=403, detail="Access denied")
        
    return {
        "id": notebook["id"],
        "name": notebook["name"],
        "description": notebook.get("description", ""),
        "color": notebook.get("color", "#1E88E5"),
        "domains": notebook.get("domains", ["General"]),
        "is_favorite": notebook.get("is_favorite", False),
        "document_count": notebook.get("document_count", 0),
        "rag_document_count": notebook.get("rag_document_count", 0),
        "created_at": notebook["created_at"].isoformat() if isinstance(notebook["created_at"], datetime.datetime) else notebook["created_at"],
        "last_accessed": notebook["last_accessed"].isoformat() if isinstance(notebook["last_accessed"], datetime.datetime) else notebook["last_accessed"]
    }

@app.put("/api/notebooks/{notebook_id}")
async def update_notebook(notebook_id: str, update: NotebookUpdate, current_user = Depends(get_current_user)):
    # Verify notebook exists and owner
    success, notebook = db.get_notebook(notebook_id)
    if not success:
        raise HTTPException(status_code=404, detail=notebook)
    if str(notebook["user_id"]) != current_user["user_id"]:
        raise HTTPException(status_code=403, detail="Access denied")
        
    data = {}
    if update.name is not None:
        data["name"] = update.name
    if update.description is not None:
        data["description"] = update.description
    if update.color is not None:
        data["color"] = update.color
    if update.domains is not None:
        data["domains"] = update.domains
        
    success, message = db.update_notebook(notebook_id, data)
    if not success:
        raise HTTPException(status_code=400, detail=message)
    return {"message": message}

@app.post("/api/notebooks/{notebook_id}/favorite")
async def toggle_favorite(notebook_id: str, current_user = Depends(get_current_user)):
    # Verify notebook
    success, notebook = db.get_notebook(notebook_id)
    if not success:
        raise HTTPException(status_code=404, detail=notebook)
    if str(notebook["user_id"]) != current_user["user_id"]:
        raise HTTPException(status_code=403, detail="Access denied")
        
    success, status = db.toggle_favorite_notebook(notebook_id)
    if not success:
        raise HTTPException(status_code=400, detail=status)
    return {"is_favorite": status}

@app.delete("/api/notebooks/{notebook_id}")
async def delete_notebook(notebook_id: str, current_user = Depends(get_current_user)):
    success, message = db.delete_notebook(notebook_id, current_user["user_id"])
    if not success:
        raise HTTPException(status_code=400, detail=message)
    return {"message": message}

# Documents

@app.get("/api/notebooks/{notebook_id}/documents")
async def list_documents(notebook_id: str, current_user = Depends(get_current_user)):
    resolved_id = get_or_create_global_notebook(current_user["user_id"]) if notebook_id == "global" else notebook_id
    success, documents = db.list_user_documents(current_user["user_id"], resolved_id)
    if not success:
        raise HTTPException(status_code=400, detail=documents)
        
    formatted = []
    for doc in documents:
        formatted.append({
            "id": doc["file_id"],
            "filename": doc["filename"],
            "display_name": doc.get("display_name", doc["filename"]),
            "file_type": doc.get("file_type", "unknown"),
            "upload_date": doc["upload_date"].isoformat() if isinstance(doc["upload_date"], datetime.datetime) else doc["upload_date"]
        })
    return formatted

@app.post("/api/notebooks/{notebook_id}/documents")
async def upload_documents(
    notebook_id: str,
    files: List[UploadFile] = File(...),
    custom_name: Optional[str] = Form(None),
    use_rag: bool = Form(True),
    llm_model: str = Form("llama-3.3-70b-versatile"),
    embedding_model: str = Form("sentence-transformers/all-MiniLM-L6-v2"),
    chunk_size: int = Form(1000),
    chunk_overlap: int = Form(200),
    use_gpu: bool = Form(True),
    split_strategy: str = Form("semantic"),
    current_user = Depends(get_current_user)
):
    user_id = current_user["user_id"]
    
    resolved_id = get_or_create_global_notebook(user_id) if notebook_id == "global" else notebook_id
    
    # Verify notebook exists and belongs to user
    success, notebook = db.get_notebook(resolved_id)
    if not success:
        raise HTTPException(status_code=404, detail=notebook)
    if str(notebook["user_id"]) != user_id:
        raise HTTPException(status_code=403, detail="Access denied")
        
    upload_success = False
    files_data = []
    
    for file in files:
        file_bytes = await file.read()
        file_type = file.filename.split('.')[-1].lower() if '.' in file.filename else 'unknown'
        
        if file_type not in ["pdf", "docx", "doc", "txt", "png", "jpg", "jpeg"]:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: .{file_type}")
            
        display_name = custom_name if custom_name and len(files) == 1 else file.filename
        
        success, result = db.save_document_file(
            file_bytes,
            file.filename,
            file_type,
            user_id,
            resolved_id,
            display_name
        )
        if success:
            upload_success = True
            files_data.append({
                "filename": file.filename,
                "content": file_bytes
            })
        else:
            print(f"Failed to upload {file.filename}: {result}")
            
    if not upload_success:
        raise HTTPException(status_code=400, detail="Failed to upload documents")
        
    metrics = None
    if use_rag and files_data:
        try:
            # Reconstruct adapted files for RAG processing
            adapted_files = []
            for fd in files_data:
                adapted_files.append(UploadedFileWrapper(fd["filename"], fd["content"]))
                
            rag = EnhancedRAG(
                llm_model_name=llm_model,
                embedding_model_name=embedding_model,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                use_gpu=use_gpu
            )
            
            # Load existing index if any
            rag.load_vector_store(db, resolved_id)
            
            # Process files synchronously
            success = rag.process_files(
                files=adapted_files,
                user_id=user_id,
                mongodb=db,
                notebook_id=resolved_id,
                is_nested=False,
                split_strategy=split_strategy
            )
            
            if success:
                # Save vector store to database
                save_success = rag.save_vector_store(db, resolved_id, user_id)
                if not save_success:
                    errors_str = "; ".join(rag.errors) if rag.errors else "Unknown vector store save error"
                    raise HTTPException(status_code=500, detail=f"Failed to save vector store: {errors_str}")
                metrics = rag.get_performance_metrics()
            else:
                errors_str = "; ".join(rag.errors) if rag.errors else "Unknown chunking error"
                raise HTTPException(status_code=500, detail=f"Failed to process files: {errors_str}")
        except HTTPException:
            raise
        except Exception as e:
            print(f"Error in synchronous RAG indexing: {str(e)}")
            raise HTTPException(status_code=500, detail=f"RAG indexing error: {str(e)}")
            
    return {
        "message": f"Successfully uploaded {len(files_data)} files.",
        "metrics": metrics
    }

@app.post("/api/notebooks/{notebook_id}/urls")
async def upload_url(
    notebook_id: str,
    request: UrlUploadRequest,
    current_user = Depends(get_current_user)
):
    user_id = current_user["user_id"]
    
    resolved_id = get_or_create_global_notebook(user_id) if notebook_id == "global" else notebook_id
    
    # Verify notebook exists and belongs to user
    success, notebook = db.get_notebook(resolved_id)
    if not success:
        raise HTTPException(status_code=404, detail=notebook)
    if str(notebook["user_id"]) != user_id:
        raise HTTPException(status_code=403, detail="Access denied")
        
    # Scrape webpage content
    scraped = scrape_webpage(request.url)
    if not scraped["success"]:
        raise HTTPException(status_code=400, detail=f"Failed to scrape webpage: {scraped.get('error')}")
        
    filename = f"url_{uuid.uuid4().hex}.txt"
    custom_name = f"URL: {scraped['title']}"
    text_content = scraped["text"]
    file_bytes = text_content.encode('utf-8')
    
    # Save document file to database/storage
    save_success, result = db.save_document_file(
        file_bytes,
        filename,
        "txt",
        user_id,
        resolved_id,
        custom_name
    )
    
    if not save_success:
        raise HTTPException(status_code=500, detail=f"Failed to save scraped content: {result}")
        
    metrics = None
    if request.use_rag:
        try:
            # Reconstruct adapted files for RAG processing
            adapted_files = [UploadedFileWrapper(filename, file_bytes)]
                
            rag = EnhancedRAG(
                llm_model_name=request.llm_model,
                embedding_model_name=request.embedding_model,
                chunk_size=request.chunk_size,
                chunk_overlap=request.chunk_overlap,
                use_gpu=request.use_gpu
            )
            
            # Load existing index if any
            rag.load_vector_store(db, resolved_id)
            
            # Process files synchronously
            process_success = rag.process_files(
                files=adapted_files,
                user_id=user_id,
                mongodb=db,
                notebook_id=resolved_id,
                is_nested=False,
                split_strategy=request.split_strategy
            )
            
            if process_success:
                # Save vector store to database
                save_idx_success = rag.save_vector_store(db, resolved_id, user_id)
                if not save_idx_success:
                    errors_str = "; ".join(rag.errors) if rag.errors else "Unknown vector store save error"
                    raise HTTPException(status_code=500, detail=f"Failed to save vector store: {errors_str}")
                metrics = rag.get_performance_metrics()
            else:
                errors_str = "; ".join(rag.errors) if rag.errors else "Unknown chunking error"
                raise HTTPException(status_code=500, detail=f"Failed to process URL: {errors_str}")
        except HTTPException:
            raise
        except Exception as e:
            print(f"Error in synchronous RAG indexing for URL: {str(e)}")
            raise HTTPException(status_code=500, detail=f"RAG indexing error: {str(e)}")
            
    return {
        "message": "Successfully uploaded URL.",
        "metrics": metrics,
        "document": {
            "id": result,
            "filename": filename,
            "display_name": custom_name,
            "file_type": "txt",
            "upload_date": datetime.datetime.now().isoformat()
        }
    }

@app.get("/api/documents/{file_id}/view")
async def view_document(file_id: str, current_user = Depends(get_current_user)):
    success, doc = db.get_document_file(file_id)
    if not success:
        raise HTTPException(status_code=404, detail=doc)
        
    file_bytes = doc["data"]
    file_type = doc["file_type"]
    filename = doc["filename"]
    
    if file_type == "pdf":
        return Response(content=file_bytes, media_type="application/pdf")
        
    elif file_type in ["png", "jpg", "jpeg"]:
        media_type = f"image/{file_type}"
        if file_type == "jpg":
            media_type = "image/jpeg"
        return Response(content=file_bytes, media_type=media_type)
        
    elif file_type == "txt":
        try:
            text_content = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text_content = file_bytes.decode('latin-1')
        return {"text": text_content, "filename": filename, "file_type": file_type}
        
    elif file_type in ["docx", "doc"]:
        # Parse DOCX paragraphs and tables
        try:
            import docx
            doc_obj = docx.Document(io.BytesIO(file_bytes))
            paragraphs = []
            for p in doc_obj.paragraphs:
                if p.text.strip():
                    style_name = p.style.name if p.style else 'Normal'
                    paragraphs.append({
                        "text": p.text,
                        "is_heading": style_name.startswith('Heading'),
                        "heading_level": int(style_name[-1]) if style_name.startswith('Heading') and style_name[-1].isdigit() else 1
                    })
            
            tables = []
            for t_idx, table in enumerate(doc_obj.tables):
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                tables.append(rows)
                
            return {
                "filename": filename,
                "file_type": file_type,
                "paragraphs": paragraphs,
                "tables": tables
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error parsing DOCX document: {str(e)}")
            
    else:
        # Stream raw bytes as download fallback
        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type="application/octet-stream",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

@app.delete("/api/documents/{file_id}")
async def delete_document(file_id: str, background_tasks: BackgroundTasks, current_user = Depends(get_current_user)):
    # Retrieve the notebook_id associated with this document before deleting it
    notebook_id = None
    try:
        doc_check = db.client.table("documents").select("notebook_id").eq("id", file_id).execute()
        if doc_check.data:
            notebook_id = doc_check.data[0].get("notebook_id")
    except Exception as check_err:
        print(f"Error checking document notebook ID: {str(check_err)}")
        
    success, result = db.delete_document(file_id, current_user["user_id"])
    if not success:
        raise HTTPException(status_code=400, detail=result)
        
    # Queue the background rebuild task
    if notebook_id:
        background_tasks.add_task(rebuild_notebook_index_task, str(notebook_id), current_user["user_id"])
        
    return {"message": result}

# RAG Chat

@app.post("/api/chat")
async def chat(request: ChatRequest, current_user = Depends(get_current_user)):
    user_id = current_user["user_id"]
    
    # Initialize RAG System
    rag = EnhancedRAG(
        llm_model_name=request.llm_model,
        embedding_model_name=request.embedding_model,
        chunk_size=request.chunk_size,
        chunk_overlap=request.chunk_overlap,
        use_gpu=request.use_gpu
    )
    
    resolved_id = get_or_create_global_notebook(user_id) if (not request.notebook_id or request.notebook_id == "global") else request.notebook_id
    
    # Load index if notebook specified
    if resolved_id:
        success = rag.load_vector_store(db, resolved_id)
        if not success:
            # Return warning if RAG selected but no vectors found
            if request.mode != "hybrid":
                return {
                    "answer": "This notebook has no processed document index yet. Please upload files and ensure 'Process with RAG' is enabled.",
                    "sources": [],
                    "query_time": 0.0,
                    "mode": request.mode
                }
                
    import re
    # Check if original query requested paragraph form
    is_para_original = bool(re.search(r'\b(paras?|paragraphs?)\b', request.query.lower()))

    standalone_query = request.query
    if request.history:
        try:
            chat_history_str = ""
            for msg in request.history:
                role = msg.get("role", "user")
                content = msg.get("content", "")
                if isinstance(content, dict):
                    content = content.get("answer", "")
                chat_history_str += f"{role.capitalize()}: {content}\n"
                
            rephrase_prompt = f"""
            Given the following conversation history and the latest user query, rephrase the latest query to be a standalone, self-contained question that retains all necessary context from the conversation history. This standalone query will be used for a database search, so make sure it contains the key search terms.
            Do NOT answer the question. Do NOT add any preamble. Just output the standalone, rephrased query and nothing else.
            
            Conversation History:
            {chat_history_str}
            
            Latest User Query: {request.query}
            
            Standalone Query:
            """
            rephrased = rag.llm.invoke(rephrase_prompt).content.strip()
            if rephrased.startswith('"') and rephrased.endswith('"'):
                rephrased = rephrased[1:-1]
            if rephrased:
                print(f"Rephrased query: '{request.query}' -> '{rephrased}'")
                standalone_query = rephrased
        except Exception as e:
            print(f"Error rephrasing query: {str(e)}")

    # If original query requested paragraph form, ensure the standalone query carries this instruction
    if is_para_original and not bool(re.search(r'\b(paras?|paragraphs?)\b', standalone_query.lower())):
        standalone_query += " (Format the answer in paragraph form only)"

    response = rag.ask(
        standalone_query,
        mode=request.mode,
        user_id=user_id,
        mongodb=db,
        notebook_id=resolved_id
    )
    
    if isinstance(response, str):
        # Fallback to string error message format
        return {
            "answer": response,
            "sources": [],
            "query_time": 0.0,
            "mode": request.mode
        }
        
    return response

# Analytics & Settings

@app.get("/api/settings/analytics")
async def get_analytics(current_user = Depends(get_current_user)):
    success, analytics = db.get_user_analytics(current_user["user_id"])
    if not success:
        raise HTTPException(status_code=400, detail=analytics)
        
    # Serialize datetimes to ISO format
    if analytics.get("created_at"):
        analytics["created_at"] = analytics["created_at"].isoformat()
    if analytics.get("last_login"):
        analytics["last_login"] = analytics["last_login"].isoformat()
    if analytics.get("last_activity"):
        analytics["last_activity"] = analytics["last_activity"].isoformat()
        
    for nb in analytics.get("notebook_stats", []):
        if nb.get("created_at"):
            nb["created_at"] = nb["created_at"].isoformat()
        if nb.get("last_accessed"):
            nb["last_accessed"] = nb["last_accessed"].isoformat()
            
    for q in analytics.get("recent_queries", []):
        if q.get("timestamp"):
            q["timestamp"] = q["timestamp"].isoformat()
            
    return analytics

@app.get("/api/settings/diagnostics")
async def get_diagnostics(current_user = Depends(get_current_user)):
    success, notebooks = db.get_notebooks(current_user["user_id"])
    if not success:
        raise HTTPException(status_code=400, detail=notebooks)
        
    diagnostics = []
    for nb in notebooks:
        vector_success, vector_result = db.get_faiss_index(nb["id"])
        if vector_success:
            status = "Vectors stored"
            metadata = vector_result.get("metadata", {})
            doc_count = metadata.get("document_count", "Unknown")
            size_bytes = metadata.get("index_size_bytes", 0)
            size = f"{size_bytes/1024/1024:.2f} MB" if size_bytes else "Unknown"
            last_updated = vector_result.get("updated_at")
            if isinstance(last_updated, datetime.datetime):
                last_updated = last_updated.isoformat()
            elif isinstance(last_updated, str):
                last_updated = last_updated
            else:
                last_updated = "Unknown"
        else:
            status = "No vectors"
            doc_count = "-"
            size = "-"
            last_updated = "-"
            
        diagnostics.append({
            "notebook_id": nb["id"],
            "name": nb["name"],
            "status": status,
            "documents": doc_count,
            "size": size,
            "last_updated": last_updated
        })
        
    return diagnostics

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("backend:app", host="0.0.0.0", port=8000, reload=True)
